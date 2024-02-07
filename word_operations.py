from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

def add_landscape_page_to_word(doc_path):
    doc = Document(doc_path)
    # Füge einen Abschnitt (Seite) im Querformat hinzu
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    # Füge der neuen Seite einen Absatz hinzu, damit sie nicht leer ist
    paragraph = section.add_paragraph()
    paragraph.add_run("Neue Seite im Querformat")
    doc.save(doc_path)